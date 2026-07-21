"""
Fills the tokenized A.B.T. employment contract template and converts it to PDF.
Works directly on the .docx bytes (no need to pre-unzip).
"""
import datetime
import io
import os
import re
import subprocess
import tempfile
import zipfile

TOKEN_PATTERN = re.compile(r"\{\{[A-Z_]+\}\}")

MONTHLY_MARKER = "שכר עובד משרה מלאה:"
HOURLY_MARKER = "שכר לעובד שעתי:"

# worker contract: schedule-type clauses (distinct from the wage markers above)
FIXED_SCHEDULE_MARKER = "<w:t>עובד משרה מלאה:</w:t>"
SHIFT_SCHEDULE_MARKER = '<w:t xml:space="preserve">עובד משמרות: </w:t>'


def _read_docx_parts(docx_bytes: bytes) -> dict:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return {name: zf.read(name) for name in zf.namelist()}


def _write_docx(parts: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in parts.items():
            zf.writestr(name, content)
    return buf.getvalue()


def _drop_paragraph_containing(xml: str, marker: str) -> str:
    """Remove the whole <w:p ...>...</w:p> block that contains `marker`."""
    idx = xml.find(marker)
    if idx == -1:
        return xml
    start = xml.rfind("<w:p ", 0, idx)
    end = xml.find("</w:p>", idx) + len("</w:p>")
    return xml[:start] + xml[end:]


def fill_document(template_path: str, fields: dict) -> bytes:
    """
    Generic filler: replaces every {{TOKEN}} in the template with fields[TOKEN].
    Works for any of the tokenized templates (termination letter, hearing
    invitation, employment confirmation, safety declaration, ...).
    Raises ValueError if any token in the template is left unfilled.
    """
    with open(template_path, "rb") as f:
        docx_bytes = f.read()

    parts = _read_docx_parts(docx_bytes)
    xml = parts["word/document.xml"].decode("utf-8")

    for key, value in fields.items():
        xml = xml.replace("{{%s}}" % key, str(value))

    leftover = TOKEN_PATTERN.findall(xml)
    if leftover:
        raise ValueError(f"Missing values for tokens: {set(leftover)}")

    parts["word/document.xml"] = xml.encode("utf-8")
    return _write_docx(parts)


def fill_contract(template_path: str, fields: dict, pay_type: str) -> bytes:
    """
    Employment-contract-specific filler: also drops whichever of the
    monthly-salary / hourly-wage clauses does not apply.

    fields must contain: EMPLOYEE_NAME, EMPLOYEE_ID, START_DATE, STATION_NAME,
    and either HOURLY_WAGE (pay_type='hourly') or MONTHLY_SALARY (pay_type='monthly').
    Returns filled .docx bytes.
    """
    with open(template_path, "rb") as f:
        docx_bytes = f.read()

    parts = _read_docx_parts(docx_bytes)
    xml = parts["word/document.xml"].decode("utf-8")

    if pay_type == "hourly":
        xml = _drop_paragraph_containing(xml, MONTHLY_MARKER)
    elif pay_type == "monthly":
        xml = _drop_paragraph_containing(xml, HOURLY_MARKER)
    else:
        raise ValueError("pay_type must be 'hourly' or 'monthly'")

    for key, value in fields.items():
        xml = xml.replace("{{%s}}" % key, str(value))

    leftover = TOKEN_PATTERN.findall(xml)
    if leftover:
        raise ValueError(f"Missing values for tokens: {set(leftover)}")

    parts["word/document.xml"] = xml.encode("utf-8")
    return _write_docx(parts)


def fill_worker_contract(template_path: str, fields: dict, pay_type: str, schedule_type: str) -> bytes:
    """
    Worker-contract-specific filler: drops whichever wage clause (monthly/hourly)
    and whichever schedule clause (fixed/shift) does not apply.

    fields must contain: EMPLOYEE_NAME, EMPLOYEE_ID, JOB_TITLE, STATION_NAME,
    START_DATE, and either HOURLY_WAGE (pay_type='hourly') or MONTHLY_SALARY
    (pay_type='monthly'). schedule_type is 'fixed' or 'shift'.
    Returns filled .docx bytes.
    """
    with open(template_path, "rb") as f:
        docx_bytes = f.read()

    parts = _read_docx_parts(docx_bytes)
    xml = parts["word/document.xml"].decode("utf-8")

    if pay_type == "hourly":
        xml = _drop_paragraph_containing(xml, MONTHLY_MARKER)
    elif pay_type == "monthly":
        xml = _drop_paragraph_containing(xml, HOURLY_MARKER)
    else:
        raise ValueError("pay_type must be 'hourly' or 'monthly'")

    if schedule_type == "shift":
        xml = _drop_paragraph_containing(xml, FIXED_SCHEDULE_MARKER)
    elif schedule_type == "fixed":
        xml = _drop_paragraph_containing(xml, SHIFT_SCHEDULE_MARKER)
    else:
        raise ValueError("schedule_type must be 'fixed' or 'shift'")

    for key, value in fields.items():
        xml = xml.replace("{{%s}}" % key, str(value))

    leftover = TOKEN_PATTERN.findall(xml)
    if leftover:
        raise ValueError(f"Missing values for tokens: {set(leftover)}")

    parts["word/document.xml"] = xml.encode("utf-8")
    return _write_docx(parts)


_HEBREW_WEEKDAY_BY_PY_WEEKDAY = {
    6: "א",     # Sunday
    0: "ב",     # Monday
    1: "ג",     # Tuesday
    2: "ד",     # Wednesday
    3: "ה",     # Thursday
    4: "ו",     # Friday
    5: "שבת",   # Saturday
}
_WEEKDAY_LETTERS = ("א", "ב", "ג", "ד", "ה", "ו", "שבת")


def _iter_all_paragraphs(container):
    """Yield every paragraph in a Document or table cell, including nested tables."""
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def fill_incident_notice(template_path: str, fields: dict) -> bytes:
    """
    Fills the insurance incident-notice form (טופס הודעה על אירוע).

    fields must contain: BRANCH_NAME, INCIDENT_DATE (DD.MM.YYYY), INCIDENT_STREET,
    INCIDENT_NUMBER, INCIDENT_CITY, INCIDENT_DESCRIPTION, CAR_PLATE, DAMAGED_PROPERTY.

    The "היום בשבוע" (day of week) line is computed automatically from INCIDENT_DATE
    and highlighted (bold + underline) on the matching letter -- no separate question.
    """
    import docx

    document = docx.Document(template_path)

    for p in _iter_all_paragraphs(document):
        for run in p.runs:
            if "{{" in run.text:
                text = run.text
                for key, value in fields.items():
                    text = text.replace("{{%s}}" % key, str(value))
                run.text = text

    leftover = set()
    for p in _iter_all_paragraphs(document):
        leftover.update(TOKEN_PATTERN.findall(p.text))
    if leftover:
        raise ValueError(f"Missing values for tokens: {leftover}")

    target_letter = None
    incident_date_str = fields.get("INCIDENT_DATE", "")
    try:
        day, month, year = incident_date_str.split(".")
        incident_date = datetime.date(int(year), int(month), int(day))
        target_letter = _HEBREW_WEEKDAY_BY_PY_WEEKDAY[incident_date.weekday()]
    except (ValueError, KeyError):
        pass

    if target_letter:
        for p in _iter_all_paragraphs(document):
            run_texts = [r.text for r in p.runs]
            if "בשבוע" in run_texts and any(t in _WEEKDAY_LETTERS for t in run_texts):
                for run in p.runs:
                    if run.text in _WEEKDAY_LETTERS:
                        is_target = run.text == target_letter
                        run.font.bold = is_target
                        run.font.underline = is_target
                break

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def fill_hearing_invitation(template_path: str, fields: dict) -> bytes:
    """
    Fills the הזמנה לשימוע (hearing invitation) template.

    Same generic {{TOKEN}} replacement as fill_document, except REASON is
    special: it may contain several reasons (one per line, or separated by
    ";"), and each one gets its own bullet line in the output -- not just a
    single line squashed into one bullet.
    """
    import copy

    import docx
    from docx.text.paragraph import Paragraph

    document = docx.Document(template_path)

    reasons_raw = fields.get("REASON", "")
    reasons = [r.strip() for r in re.split(r"[\n;]+", reasons_raw) if r.strip()]
    if not reasons:
        reasons = [reasons_raw]

    bullet_paragraph = None
    for p in _iter_all_paragraphs(document):
        if "{{REASON}}" in p.text:
            bullet_paragraph = p
            break

    if bullet_paragraph is not None:
        anchor = bullet_paragraph._p
        parent = bullet_paragraph._parent
        for reason in reasons:
            new_p_element = copy.deepcopy(anchor)
            new_paragraph = Paragraph(new_p_element, parent)
            for run in new_paragraph.runs:
                if "{{REASON}}" in run.text:
                    run.text = run.text.replace("{{REASON}}", reason)
            anchor.addprevious(new_p_element)
        anchor.getparent().remove(anchor)

    other_fields = {k: v for k, v in fields.items() if k != "REASON"}
    for p in _iter_all_paragraphs(document):
        for run in p.runs:
            if "{{" in run.text:
                text = run.text
                for key, value in other_fields.items():
                    text = text.replace("{{%s}}" % key, str(value))
                run.text = text

    leftover = set()
    for p in _iter_all_paragraphs(document):
        leftover.update(TOKEN_PATTERN.findall(p.text))
    if leftover:
        raise ValueError(f"Missing values for tokens: {leftover}")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert docx bytes to pdf bytes via headless LibreOffice. Requires `soffice` on PATH."""
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "contract.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path],
            check=True, timeout=60,
        )
        pdf_path = os.path.join(tmp, "contract.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()


def merge_pdfs(pdf_bytes_list: list) -> bytes:
    """Concatenate several PDFs (bytes) into one, in the given order."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        writer.append(io.BytesIO(pdf_bytes))
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()
